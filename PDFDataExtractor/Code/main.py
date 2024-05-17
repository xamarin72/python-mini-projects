import os
import fitz
import re
from docx import Document
from docx.shared import Inches

# Function to extract text, candidate name, and picture
def extract_data_from_pdf(pdf_path):
    doc = fitz.open(pdf_path)
    text = ""
    for page_num in range(min(2, doc.page_count)):  # Read only the first page for efficiency
        page = doc.load_page(page_num)
        text += page.get_text("text")
    doc.close()

    # Extract name and picture
    name = re.search(r'^[A-Z\s]+$', text, re.MULTILINE)
    picture = extract_picture_from_pdf(pdf_path)  # Extract picture
    return text, name.group(0) if name else "Name not found", picture if picture else "Picture not found", os.path.basename(pdf_path)

# Function to search for keywords in text
def search_keywords_in_text(text, keywords, pdf_name):
    found_keywords = []
    for keyword in keywords:
        matches = re.finditer(r'(?=(\b' + re.escape(keyword) + r'\b))', text, re.IGNORECASE)
        for match in matches:
            line_num = text.count('\n', 0, match.start()) + 1
            start = max(0, text.rfind('\n', 0, match.start()) + 1)
            end = text.find('\n', match.end())
            line = text[start:end].strip()
            if line and line not in found_keywords:
                found_keywords.append({"keyword": keyword, "line_num": line_num, "pdf_name": pdf_name, "line": line})
    return found_keywords

# Function to extract picture from the first page of the PDF near the candidate's name
def extract_picture_from_pdf(pdf_path):
    doc = fitz.open(pdf_path)
    first_page = doc.load_page(0)  # Load the first page
    xobjects = first_page.get_xobjects()  # Get XObjects from the first page

    # Iterate through XObjects to find the image near the candidate's name
    for xref in xobjects:
        xobj = xobjects[xref]
        if xobj['/Subtype'] == '/Image':
            # Assuming the image is found near the candidate's name
            image_path = f"{pdf_path}_extracted_image.png"  # File path to save the extracted image
            img = doc.extract_image(xref)  # Extract the image
            with open(image_path, "wb") as img_file:
                img_file.write(img["image"])
            return image_path  # Return the file path of the extracted image
    return None  # Return None if no image is found

# Main function to process all PDFs in a folder and save results in the 'Output' folder
def process_pdfs_in_folder(folder_path, keywords):
    output_folder = '../Output'
    os.makedirs(output_folder, exist_ok=True)  # Create the 'Output' folder if it doesn't exist
    result_doc_path = os.path.join(output_folder, 'CV_Search_Results.docx')  # Path for the output Word document

    result_doc = Document()
    for filename in os.listdir(folder_path):
        if filename.endswith(".pdf"):
            pdf_path = os.path.join(folder_path, filename)
            extracted_text, name, picture, pdf_name = extract_data_from_pdf(pdf_path)
            keyword_matches = search_keywords_in_text(extracted_text, keywords, pdf_name)


# Add results to the document
            result_doc.add_heading(f"Candidate: {name}", level=1)
            result_doc.add_paragraph(f"Document: {pdf_name}")
            if picture != "Picture not found":
                result_doc.add_picture(picture, width=Inches(2))
            # Add results to the document
            for match in keyword_matches:
                #result_doc.add_paragraph(f"-{match['keyword']} :")
                result_doc.add_paragraph(f"{match['line_num']}: {match['line']}")
            result_doc.add_page_break()

    result_doc.save(result_doc_path)  # Save the output Word document in the 'Output' folder

# Input keywords to search
keywords_input = input("Enter keywords to search (separated by commas): ")
keywords = [keyword.strip() for keyword in keywords_input.split(',')]

# Input folder path
folder_path = "../PDFs"

# Process PDFs in the folder
process_pdfs_in_folder(folder_path, keywords)
